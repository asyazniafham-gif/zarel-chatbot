"""
RAG Provider — Zarel Sevan Chatbot
ChromaDB + Mesolitica Embeddings + DeepSeek
CPU mode for AMD GPU (no CUDA)
"""

import os
import re
import json
import time
import hashlib
import asyncio
import sqlite3
import threading
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Tuple, Generator
from datetime import datetime

import torch

DEVICE = os.environ.get("EMBEDDING_DEVICE", "cuda" if torch.cuda.is_available() else "cpu")
CROSS_ENCODER_DEVICE = os.environ.get("CROSS_ENCODER_DEVICE", DEVICE)
EMBEDDING_BATCH_SIZE = int(os.environ.get("EMBEDDING_BATCH_SIZE", "256"))
INGEST_BATCH_SIZE = int(os.environ.get("INGEST_BATCH_SIZE", "500"))
print(f"[ZS] Compute device: {DEVICE} | Cross-encoder: {CROSS_ENCODER_DEVICE} | "
      f"Embedding batch: {EMBEDDING_BATCH_SIZE} | Ingest batch: {INGEST_BATCH_SIZE}")


# ============================================================================
# BASE INTERFACES
# ============================================================================
class BaseRetriever(ABC):
    @abstractmethod
    async def retrieve(self, query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
        pass

class BaseReranker(ABC):
    @abstractmethod
    def rerank(self, query: str, documents: List[Dict[str, Any]], top_n: int = 5) -> List[Dict[str, Any]]:
        pass

class BaseGenerator(ABC):
    @abstractmethod
    def generate(self, prompt: str) -> str:
        pass

    @abstractmethod
    def generate_stream(self, prompt: str) -> Generator[str, None, None]:
        pass


# ============================================================================
# EMBEDDING MODEL (Mesolitica Mistral 191M — best Malay embeddings)
# ============================================================================
_embedding_model = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer
        model_name = os.environ.get(
            "LOCAL_EMBEDDING_MODEL",
            "mesolitica/mistral-embedding-191m-8k-contrastive"
        )
        print(f"[ZS] Loading embedding model: {model_name} -> {DEVICE}")
        _embedding_model = SentenceTransformer(model_name, device=DEVICE)
        _ = _embedding_model.encode(["warmup"], device=DEVICE, batch_size=1)
        print(f"[ZS] Embedding model loaded on {DEVICE}")
    return _embedding_model


# ============================================================================
# CHROMADB
# ============================================================================
_chroma_collection = None

def get_chroma_collection():
    global _chroma_collection
    if _chroma_collection is None:
        import chromadb
        from agency_config import CHROMA_COLLECTION_NAME

        persist_dir = os.environ.get("CHROMA_PERSIST_DIR", "/app/chroma_db")
        print(f"[ZS] Initialising ChromaDB at: {persist_dir}")
        client = chromadb.PersistentClient(path=persist_dir)
        _chroma_collection = client.get_or_create_collection(
            name=CHROMA_COLLECTION_NAME,
            metadata={
                "hnsw:space": "cosine",
                "hnsw:M": 32,
                "hnsw:construction_ef": 200,
                "hnsw:search_ef": 100,
                "hnsw:num_threads": 8,
            }
        )
        print(f"[ZS] ChromaDB collection ready: {_chroma_collection.count()} documents")
    return _chroma_collection


# ============================================================================
# DOCUMENT EXTRACTION
# ============================================================================
def extract_text_from_pdf(filepath: str) -> str:
    try:
        import fitz
        doc = fitz.open(filepath)
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"[Halaman {page_num + 1}]\n{text}")
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"[ZS] PDF extraction error for {filepath}: {e}")
        return ""


def extract_text_from_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"[ZS] DOCX extraction error for {filepath}: {e}")
        return ""


def extract_text_from_file(filepath: str) -> str:
    ext = filepath.rsplit('.', 1)[-1].lower() if '.' in filepath else ''
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext == 'docx':
        return extract_text_from_docx(filepath)
    elif ext in ('md', 'txt'):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            print(f"[ZS] Text read error for {filepath}: {e}")
            return ""
    return ""


# ============================================================================
# KNOWLEDGE INGESTION
# ============================================================================
_local_docs_ingested = False

def ingest_knowledge_to_chroma(knowledge_path: str = None):
    global _local_docs_ingested
    if _local_docs_ingested:
        return

    import glob as globlib
    from agency_config import KNOWLEDGE_DIR, DOCUMENTS_DIR, CHROMA_DB_DIR

    if knowledge_path is None:
        knowledge_path = os.environ.get("LOCAL_KNOWLEDGE_PATH", KNOWLEDGE_DIR)

    collection = get_chroma_collection()

    if collection.count() > 0:
        # Check if any knowledge file is newer than the last ingest timestamp
        stamp_file = os.path.join(os.environ.get("CHROMA_PERSIST_DIR", CHROMA_DB_DIR), ".last_ingest")
        last_ingest = os.path.getmtime(stamp_file) if os.path.exists(stamp_file) else 0
        all_files = []
        for ext in ('*.md', '*.txt', '*.pdf', '*.docx'):
            all_files += globlib.glob(os.path.join(knowledge_path, ext))
        newest = max((os.path.getmtime(f) for f in all_files), default=0)
        if newest <= last_ingest:
            print(f"[ZS] ChromaDB up to date ({collection.count()} chunks), skipping ingestion")
            _local_docs_ingested = True
            return
        print(f"[ZS] Knowledge files changed — re-ingesting...")
        collection.delete(where={"source": {"$ne": ""}})

    print(f"[ZS] Starting knowledge ingestion...")

    model = get_embedding_model()

    # Gather all supported files
    all_files = []
    for ext in ('*.md', '*.txt', '*.pdf', '*.docx'):
        all_files += globlib.glob(os.path.join(knowledge_path, ext))

    local_docs_path = os.environ.get("LOCAL_DOCUMENTS_PATH", DOCUMENTS_DIR)
    if os.path.exists(local_docs_path):
        for ext in ('*.md', '*.txt', '*.pdf', '*.docx'):
            all_files += globlib.glob(os.path.join(local_docs_path, ext))

    print(f"[ZS] Found {len(all_files)} files to ingest")

    chunk_size = 600  # slightly larger for technical docs
    chunk_overlap = 120

    all_ids = []
    all_texts = []
    all_metadatas = []
    skipped = 0

    for filepath in all_files:
        try:
            content = extract_text_from_file(filepath)
            if not content or len(content) < 50:
                skipped += 1
                continue

            filename = os.path.basename(filepath)
            file_ext = filepath.rsplit('.', 1)[-1].lower()

            if file_ext == 'pdf':
                doc_type = "Dokumen PDF"
            elif file_ext == 'docx':
                doc_type = "Dokumen DOCX"
            else:
                doc_type = "Fail Pengetahuan"

            # Build page offset map for PDFs (find all [Halaman N] markers)
            page_markers = []
            if file_ext == 'pdf':
                for m in re.finditer(r'\[Halaman (\d+)\]', content):
                    page_markers.append((m.start(), int(m.group(1))))

            def _get_pages_for_range(start: int, end: int) -> str:
                """Return page number(s) that a chunk spans, e.g. '13' or '13-14'."""
                if not page_markers:
                    return ""
                # Find the page active at chunk start (last marker before start)
                start_page = None
                for pos, page_num in page_markers:
                    if pos <= start:
                        start_page = page_num
                    else:
                        break
                # Find any page transitions within the chunk
                pages = set()
                if start_page:
                    pages.add(start_page)
                for pos, page_num in page_markers:
                    if pos >= end:
                        break
                    if pos > start:
                        pages.add(page_num)
                if not pages:
                    return ""
                sorted_pages = sorted(pages)
                if len(sorted_pages) == 1:
                    return str(sorted_pages[0])
                return f"{sorted_pages[0]}-{sorted_pages[-1]}"

            # Split into chunks
            chunks = []
            chunk_pages = []
            step = chunk_size - chunk_overlap
            for i in range(0, len(content), step):
                chunk = content[i:i + chunk_size]
                if len(chunk.strip()) > 50:
                    chunks.append(chunk)
                    chunk_pages.append(_get_pages_for_range(i, i + len(chunk)))

            for idx, chunk in enumerate(chunks):
                doc_id = hashlib.md5(f"{filename}:{idx}".encode()).hexdigest()
                all_ids.append(doc_id)
                all_texts.append(chunk)
                meta = {
                    "filename": filename,
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "source": filepath,
                    "type": doc_type,
                }
                if chunk_pages[idx]:
                    meta["page"] = chunk_pages[idx]
                all_metadatas.append(meta)

            print(f"[ZS] Chunked {filename}: {len(chunks)} chunks ({len(content)} chars extracted)")

        except Exception as e:
            print(f"[ZS] Error processing {filepath}: {e}")
            skipped += 1

    if skipped:
        print(f"[ZS] Skipped {skipped} files (empty or unreadable)")

    if all_texts:
        print(f"[ZS] Embedding {len(all_texts)} chunks on {DEVICE} "
              f"(batch_size={EMBEDDING_BATCH_SIZE})...")
        t0 = time.time()
        all_embeddings = model.encode(
            all_texts,
            batch_size=EMBEDDING_BATCH_SIZE,
            show_progress_bar=True,
            device=DEVICE,
            normalize_embeddings=True,
        ).tolist()
        embed_time = time.time() - t0
        print(f"[ZS] Embedded {len(all_texts)} chunks in {embed_time:.1f}s "
              f"({len(all_texts)/embed_time:.0f} chunks/s)")

        for i in range(0, len(all_ids), INGEST_BATCH_SIZE):
            end = min(i + INGEST_BATCH_SIZE, len(all_ids))
            collection.add(
                ids=all_ids[i:end],
                documents=all_texts[i:end],
                metadatas=all_metadatas[i:end],
                embeddings=all_embeddings[i:end]
            )

        total = len(all_ids)
        print(f"[ZS] Ingested {total} chunks into ChromaDB")
        # Write timestamp so next check knows when last ingest occurred
        stamp_file = os.path.join(os.environ.get("CHROMA_PERSIST_DIR", CHROMA_DB_DIR), ".last_ingest")
        with open(stamp_file, 'w') as f:
            f.write(str(time.time()))

    _local_docs_ingested = True


# ============================================================================
# RETRIEVER — Hybrid search + LLM Query Expansion
# ============================================================================
class HybridRetriever(BaseRetriever):
    """
    Ultra retriever with:
    1. Query expansion (LLM generates 3 search variants)
    2. Hybrid search (vector + BM25 keyword)
    3. Multi-query fusion via RRF
    """

    def __init__(self, generator: BaseGenerator = None):
        self.top_k = int(os.environ.get("RETRIEVAL_TOP_K", "15"))
        self.generator = generator

    def expand_query(self, query: str) -> List[str]:
        """Use LLM to generate query variants for better recall."""
        if not self.generator:
            return [query]

        prompt = """You are a document search expert.
Write 3 different versions of the following question for document retrieval. Use relevant technical terms.

Original question: """ + query + """

Reply with ONLY 3 lines, one question per line. No numbers or bullets:"""

        try:
            result = self.generator.generate(prompt)
            variants = [v.strip() for v in result.strip().split('\n') if v.strip()]
            all_queries = [query] + variants[:3]
            print(f"[ZS] Query expansion: {query[:60]} -> {len(all_queries)} variants")
            for i, q in enumerate(all_queries):
                print(f"  [{i}] {q[:80]}")
            return all_queries
        except Exception as e:
            print(f"[ZS] Query expansion failed: {e}")
            return [query]

    # Cached BM25 index — rebuilt only when collection changes
    _bm25_cache = None
    _bm25_doc_count = 0

    @classmethod
    def _get_bm25(cls):
        """Get or build cached BM25 index. Only rebuilds if doc count changes."""
        collection = get_chroma_collection()
        current_count = collection.count()
        if cls._bm25_cache is None or cls._bm25_doc_count != current_count:
            from rank_bm25 import BM25Okapi
            all_docs = collection.get(include=["documents", "metadatas"])
            if all_docs and all_docs['documents']:
                tokenized = [d.lower().split() for d in all_docs['documents']]
                cls._bm25_cache = (BM25Okapi(tokenized), all_docs)
                cls._bm25_doc_count = current_count
                print(f"[ZS] BM25 index built: {current_count} docs (cached)")
        return cls._bm25_cache

    def _bm25_search(self, query: str) -> Dict:
        """BM25 keyword search using cached index."""
        bm25_scores = {}
        try:
            cache = self._get_bm25()
            if cache:
                bm25, all_docs = cache
                query_tokens = query.lower().split()
                scores = bm25.get_scores(query_tokens)
                for i, (doc, meta, score) in enumerate(zip(
                    all_docs['documents'], all_docs['metadatas'], scores
                )):
                    if score > 0:
                        doc_id = hashlib.md5(doc[:200].encode()).hexdigest()
                        bm25_scores[doc_id] = (doc, meta, float(score))
            print(f"[ZS] BM25 found {len(bm25_scores)} keyword matches")
        except Exception as e:
            print(f"[ZS] BM25 search failed: {e}")
        return bm25_scores

    async def retrieve(self, query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
        await asyncio.to_thread(ingest_knowledge_to_chroma)

        collection = get_chroma_collection()
        model = get_embedding_model()

        # Run query expansion + BM25 IN PARALLEL (BM25 doesn't need expanded queries)
        expand_task = asyncio.to_thread(self.expand_query, query)
        bm25_task = asyncio.to_thread(self._bm25_search, query)
        expanded, bm25_scores = await asyncio.gather(expand_task, bm25_task)

        # Batch-encode all query variants on GPU
        all_embeddings = model.encode(
            expanded, batch_size=len(expanded), device=DEVICE, normalize_embeddings=True
        )

        # Vector search across all variants
        vector_results = {}
        for q, q_embedding in zip(expanded, all_embeddings):
            results = collection.query(
                query_embeddings=[q_embedding.tolist()],
                n_results=self.top_k,
                include=["documents", "metadatas", "distances"]
            )
            if results and results['documents']:
                for doc, meta, dist in zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                ):
                    doc_id = hashlib.md5(doc[:200].encode()).hexdigest()
                    score = 1.0 - dist
                    if doc_id not in vector_results or score > vector_results[doc_id][2]:
                        vector_results[doc_id] = (doc, meta, score)

        # RRF fusion
        fused = {}

        vector_sorted = sorted(vector_results.items(), key=lambda x: x[1][2], reverse=True)
        for rank, (doc_id, (doc, meta, score)) in enumerate(vector_sorted):
            rrf_score = 1.0 / (60 + rank)
            fused[doc_id] = {
                "doc": doc, "meta": meta,
                "vector_score": score, "bm25_score": 0,
                "rrf_score": rrf_score
            }

        bm25_sorted = sorted(bm25_scores.items(), key=lambda x: x[1][2], reverse=True)
        for rank, (doc_id, (doc, meta, score)) in enumerate(bm25_sorted):
            rrf_addition = 1.0 / (60 + rank)
            if doc_id in fused:
                fused[doc_id]["bm25_score"] = score
                fused[doc_id]["rrf_score"] += rrf_addition
            else:
                fused[doc_id] = {
                    "doc": doc, "meta": meta,
                    "vector_score": 0, "bm25_score": score,
                    "rrf_score": rrf_addition
                }

        final_sorted = sorted(fused.items(), key=lambda x: x[1]["rrf_score"], reverse=True)

        contexts = []
        sources = []
        for doc_id, data in final_sorted[:self.top_k]:
            filename = data["meta"].get("filename", "Unknown")
            page = data["meta"].get("page", "")
            page_label = f" (Halaman {page})" if page else ""
            contexts.append(f"[SUMBER - {filename}{page_label}]\n{data['doc']}")
            source_entry = {
                "type": "Dokumen (Hybrid Search)",
                "filename": filename,
                "page_content": data["doc"][:500],
                "score": round(data["rrf_score"], 4),
                "vector_score": round(data["vector_score"], 4),
                "bm25_score": round(data["bm25_score"], 4),
                "source_uri": f"local://{filename}",
                "chunk_index": data["meta"].get("chunk_index", 0),
                "priority": "PRIMARY",
            }
            if page:
                source_entry["page"] = page
            sources.append(source_entry)

        print(f"[ZS] Hybrid retrieval: {len(contexts)} results "
              f"(vector: {len(vector_results)}, BM25: {len(bm25_scores)}, "
              f"fused: {len(fused)}, queries: {len(expanded)})")

        return contexts, sources


# ============================================================================
# RERANKER — Cross-encoder on GPU
# ============================================================================
class CrossEncoderReranker(BaseReranker):
    def __init__(self):
        self._cross_encoder = None

    def _get_model(self):
        if self._cross_encoder is None:
            from sentence_transformers import CrossEncoder
            model_name = os.environ.get(
                "RERANKER_MODEL",
                "cross-encoder/ms-marco-MiniLM-L-6-v2"
            )
            print(f"[ZS] Loading cross-encoder: {model_name} -> {CROSS_ENCODER_DEVICE}")
            self._cross_encoder = CrossEncoder(model_name, device=CROSS_ENCODER_DEVICE)
            _ = self._cross_encoder.predict([("warmup query", "warmup doc")])
            print(f"[ZS] Cross-encoder ready on {CROSS_ENCODER_DEVICE}")
        return self._cross_encoder

    def rerank(self, query: str, documents: List[Dict[str, Any]], top_n: int = 5) -> List[Dict[str, Any]]:
        if not documents or len(documents) <= top_n:
            return documents

        try:
            model = self._get_model()
            pairs = [
                (query, doc.get("page_content", "")[:512])
                for doc in documents
            ]
            scores = model.predict(pairs, batch_size=len(pairs))

            scored = list(zip(scores, documents))
            scored.sort(key=lambda x: x[0], reverse=True)

            result = []
            for score, doc in scored[:top_n]:
                doc_copy = doc.copy()
                doc_copy["rerank_score"] = float(score)
                result.append(doc_copy)

            print(f"[ZS] Cross-encoder reranking: {len(documents)} -> {len(result)} docs")
            return result

        except Exception as e:
            print(f"[ZS] Cross-encoder reranking failed: {e}")
            return documents[:top_n]


# ============================================================================
# LLM GENERATOR — DeepSeek (OpenAI-compatible API)
# ============================================================================
class DeepSeekGenerator(BaseGenerator):
    """DeepSeek API generator via OpenAI-compatible endpoint."""

    MAX_RETRIES = 2
    TIMEOUT = 90

    def __init__(self, model: str = None):
        self.api_key = os.environ.get("DEEPSEEK_API_KEY", "")
        if not self.api_key:
            raise ValueError("DEEPSEEK_API_KEY not set — required for Zarel Sevan Chatbot")
        self.base_url = os.environ.get("DEEPSEEK_API_BASE", "https://api.deepseek.com/v1")
        self.model = model or os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")
        self._client = None

    def _get_client(self):
        if self._client is None:
            from openai import OpenAI
            self._client = OpenAI(
                api_key=self.api_key,
                base_url=self.base_url,
                timeout=self.TIMEOUT,
                max_retries=0,
            )
        return self._client

    def generate(self, prompt: str) -> str:
        last_err = None
        for attempt in range(self.MAX_RETRIES + 1):
            try:
                client = self._get_client()
                response = client.chat.completions.create(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=4000,
                    temperature=0.2,
                )
                result = response.choices[0].message.content or ""
                if result:
                    return result
                print(f"[DeepSeek] Warning: empty response on attempt {attempt+1}")
            except Exception as e:
                last_err = e
                print(f"[DeepSeek] Error attempt {attempt+1}/{self.MAX_RETRIES+1}: {e}")
                if attempt < self.MAX_RETRIES:
                    time.sleep(1)
        print(f"[DeepSeek] All retries failed: {last_err}")
        return "Sorry, the system is currently busy. Please try again."

    def generate_stream(self, prompt: str) -> Generator[str, None, None]:
        try:
            client = self._get_client()
            stream = client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=0.2,
                stream=True,
            )
            got_content = False
            for chunk in stream:
                if chunk.choices and chunk.choices[0].delta.content:
                    got_content = True
                    yield chunk.choices[0].delta.content
            if not got_content:
                print(f"[DeepSeek] Stream produced no content, falling back to non-stream")
                yield self.generate(prompt)
        except Exception as e:
            print(f"[DeepSeek] Stream error: {e}")
            yield "Sorry, an error occurred. Please try again."


# ============================================================================
# CONVERSATION MEMORY (SQLite)
# ============================================================================
class ConversationMemory:
    """SQLite-backed conversation memory for multi-turn context."""

    def __init__(self, db_path: str = "/app/logs/memory.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        conn = sqlite3.connect(self.db_path)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT,
                role TEXT,
                content TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.commit()
        conn.close()

    def save_turn(self, session_id: str, role: str, content: str):
        conn = sqlite3.connect(self.db_path)
        conn.execute(
            "INSERT INTO conversations (session_id, role, content) VALUES (?, ?, ?)",
            (session_id, role, content)
        )
        conn.commit()
        conn.close()

    def get_history(self, session_id: str, limit: int = 10) -> List[Dict[str, str]]:
        conn = sqlite3.connect(self.db_path)
        rows = conn.execute(
            "SELECT role, content FROM conversations WHERE session_id = ? ORDER BY id DESC LIMIT ?",
            (session_id, limit * 2)
        ).fetchall()
        conn.close()
        return [{"role": r, "content": c} for r, c in reversed(rows)]


# ============================================================================
# ULTRA ENHANCER — Self-eval, follow-ups, enhanced prompt
# ============================================================================
class UltraEnhancer:
    """
    Ultra mode enhancements:
    - Self-evaluation (3-dimension scoring)
    - Follow-up question suggestions
    - Enhanced chain-of-thought prompt
    """

    def __init__(self, generator: BaseGenerator):
        self.generator = generator
        self.memory = ConversationMemory()

    def self_evaluate(self, query: str, answer: str, contexts: List[str]) -> Dict[str, Any]:
        """LLM evaluates its own answer quality on 3 dimensions."""
        context_preview = "\n".join(c[:200] for c in contexts[:3])

        eval_prompt = f"""Anda pemeriksa kualiti jawapan AI. Nilai jawapan berikut:

SOALAN: {query}

KONTEKS YANG ADA:
{context_preview}

JAWAPAN AI:
{answer[:1000]}

Nilai 1-5 untuk setiap kriteria:
1. RELEVAN: Adakah jawapan menjawab soalan? (1=tidak relevan, 5=sangat relevan)
2. TEPAT: Adakah jawapan berdasarkan dokumen? (1=reka sendiri, 5=berdasarkan sumber)
3. LENGKAP: Adakah jawapan cukup lengkap? (1=tidak lengkap, 5=sangat lengkap)

Jawab HANYA dalam format JSON:
{{"relevan": N, "tepat": N, "lengkap": N, "purata": N.N, "nota": "satu ayat ulasan"}}"""

        try:
            result = self.generator.generate(eval_prompt)
            json_match = re.search(r'\{[^}]+\}', result)
            if json_match:
                eval_data = json.loads(json_match.group())
                if "purata" not in eval_data:
                    scores = [eval_data.get("relevan", 3), eval_data.get("tepat", 3), eval_data.get("lengkap", 3)]
                    eval_data["purata"] = round(sum(scores) / len(scores), 1)
                print(f"[ZS] Self-eval: {eval_data.get('purata', '?')}/5 — {eval_data.get('nota', '')}")
                return eval_data
        except Exception as e:
            print(f"[ZS] Self-evaluation failed: {e}")

        return {"relevan": 3, "tepat": 3, "lengkap": 3, "purata": 3.0, "nota": "Penilaian automatik gagal"}

    def suggest_followups(self, query: str, answer: str) -> List[str]:
        """Generate follow-up question suggestions."""
        prompt = f"""Based on this conversation, suggest 3 useful follow-up questions in English.

User question: {query}
Answer summary: {answer[:500]}

Write 3 follow-up questions in English. One question per line. No numbers:"""

        try:
            result = self.generator.generate(prompt)
            suggestions = [s.strip() for s in result.strip().split('\n') if s.strip() and len(s.strip()) > 10]
            return suggestions[:3]
        except Exception as e:
            print(f"[ZS] Follow-up suggestion failed: {e}")
            return []


# ============================================================================
# WEB SEARCH — Tavily + Brave (dual provider)
# ============================================================================
_tavily_client = None
_web_search_lock = threading.Lock()

def _get_tavily_client():
    global _tavily_client
    with _web_search_lock:
        if _tavily_client is None:
            key = os.environ.get("TAVILY_API_KEY", "")
            if key:
                try:
                    from tavily import TavilyClient
                    _tavily_client = TavilyClient(api_key=key)
                    print("[ZS] Tavily web search client initialised")
                except ImportError:
                    print("[ZS] tavily-python not installed, Tavily search disabled")
    return _tavily_client

from agency_config import WEB_SEARCH_PREFIX

def search_web_tavily(query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Search the web using Tavily, focused on PT topics."""
    try:
        client = _get_tavily_client()
        if not client:
            return [], []

        q_lower = query.lower()
        if not any(kw in q_lower for kw in ("zarel", "sevan", "produk", "tiktok")):
            search_query = f"{WEB_SEARCH_PREFIX} {query}"
        else:
            search_query = query

        response = client.search(
            query=search_query,
            search_depth="advanced",
            max_results=5,
            include_answer=False,
            include_raw_content=False,
        )

        contexts, sources = [], []
        for result in response.get("results", [])[:5]:
            title = result.get("title", "")
            url = result.get("url", "")
            content = result.get("content", "")
            if title and content:
                contexts.append(f"[SUMBER WEB] {title}\nURL: {url}\n{content}")
                sources.append({
                    "type": "Web (Tavily)",
                    "filename": title,
                    "page_content": content[:500],
                    "score": result.get("score", 0),
                    "source_uri": url,
                    "priority": "SUPPLEMENTARY",
                })
        print(f"[ZS] Tavily search: {len(contexts)} results for: {search_query[:80]}")
        return contexts, sources
    except Exception as e:
        print(f"[ZS] Tavily search error: {e}")
        return [], []


def search_web_brave(query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Search the web using Brave Search API, focused on PT topics."""
    brave_key = os.environ.get("BRAVE_API_KEY", "")
    if not brave_key:
        return [], []
    try:
        import httpx

        q_lower = query.lower()
        if not any(kw in q_lower for kw in ("zarel", "sevan", "produk", "tiktok")):
            search_query = f"{WEB_SEARCH_PREFIX} {query}"
        else:
            search_query = query

        resp = httpx.get(
            "https://api.search.brave.com/res/v1/web/search",
            params={"q": search_query, "count": 5},
            headers={"X-Subscription-Token": brave_key, "Accept": "application/json"},
            timeout=15.0,
        )
        resp.raise_for_status()
        data = resp.json()

        contexts, sources = [], []
        for result in data.get("web", {}).get("results", [])[:5]:
            title = result.get("title", "")
            url = result.get("url", "")
            description = result.get("description", "")
            if title and description:
                contexts.append(f"[SUMBER WEB] {title}\nURL: {url}\n{description}")
                sources.append({
                    "type": "Web (Brave)",
                    "filename": title,
                    "page_content": description[:500],
                    "score": 0,
                    "source_uri": url,
                    "priority": "SUPPLEMENTARY",
                })
        print(f"[ZS] Brave search: {len(contexts)} results for: {search_query[:80]}")
        return contexts, sources
    except Exception as e:
        print(f"[ZS] Brave search error: {e}")
        return [], []


def search_web(query: str) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Dual-provider web search: Tavily -> Brave -> empty."""
    if os.environ.get("TAVILY_API_KEY"):
        contexts, sources = search_web_tavily(query)
        if contexts:
            return contexts, sources

    if os.environ.get("BRAVE_API_KEY"):
        contexts, sources = search_web_brave(query)
        if contexts:
            return contexts, sources

    return [], []


# ============================================================================
# PROVIDER FACTORY
# ============================================================================
_providers = None

def get_providers() -> Dict[str, Any]:
    global _providers
    if _providers is None:
        print(f"\n{'='*60}")
        print(f"  INITIALISING ZAREL SEVAN CHATBOT PROVIDERS")
        print(f"{'='*60}\n")

        # Single Claude model for all tasks
        model_main = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")
        generator = DeepSeekGenerator(model=model_main)
        generator_fast = generator  # reuse same model for utilities

        enhancer = UltraEnhancer(generator_fast)

        _providers = {
            "mode": "ultra",
            "retriever": HybridRetriever(generator=generator_fast),
            "reranker": CrossEncoderReranker(),
            "generator": generator,
            "generator_fast": generator_fast,
            "enhancer": enhancer,
            "memory": enhancer.memory,
            "web_search": search_web,
            "description": f"Zarel Sevan Chatbot ({model_main} + web search)"
        }

        web_providers = []
        if os.environ.get("TAVILY_API_KEY"):
            web_providers.append("Tavily")
        if os.environ.get("BRAVE_API_KEY"):
            web_providers.append("Brave")
        web_str = " + ".join(web_providers) if web_providers else "disabled"

        print(f"[ZS] LLM: {model_main}")
        print(f"[ZS] Web search: {web_str}")
        print(f"[ZS] Providers initialised: {_providers['description']}")

    return _providers


def get_mode_info() -> Dict[str, Any]:
    providers = get_providers()
    model = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")
    return {
        "current_mode": "ultra",
        "description": providers["description"],
        "features": [
            "Query Expansion (3 LLM variants)",
            "Hybrid Search (Vector + BM25 + RRF)",
            "Cross-Encoder Reranking (CPU)",
            "Self-Evaluation (3-dimension scoring)",
            "Follow-up Suggestions",
            "Conversation Memory (SQLite)",
            "Chain-of-Thought Prompting",
            "Web Search (Tavily + Brave)",
        ],
        "llm": f"DeepSeek ({model})",
        "embeddings": "mesolitica/mistral-embedding-191m-8k-contrastive (Malay, CPU)",
        "reranker": "cross-encoder/ms-marco-MiniLM-L-6-v2 (CPU)",
        "vector_db": "ChromaDB (hybrid: vector + BM25 + query expansion)",
        "device": DEVICE,
    }
